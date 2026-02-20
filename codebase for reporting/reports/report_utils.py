import io
import csv
import ast
from datetime import datetime
from typing import List, Sequence, Tuple, Dict, Any
import os
from pathlib import Path

from docx import Document

from .pdf_utils import render_report_html


# -------------------------
# Column normalisation
# -------------------------

def normalize_display_columns(columns: List[str]) -> Tuple[List[str], Dict[str, str]]:
    """
    Given a list of raw column names (e.g. ['ORDER_NO', 'CLIENT_NAME']),
    return:
        - display_columns: ['Order No', 'Client Name']
        - column_display_map: {'ORDER_NO': 'Order No', 'CLIENT_NAME': 'Client Name'}

    Rules:
      - strip whitespace
      - replace underscores with spaces
      - collapse multiple spaces
      - title case
    """
    display_cols: List[str] = []
    mapping: Dict[str, str] = {}

    for raw in columns:
        raw_str = "" if raw is None else str(raw)
        base = raw_str.strip().replace("_", " ")
        # collapse multiple spaces
        base = " ".join(base.split())
        if base:
            display = base.lower().title()
        else:
            # fallback: keep raw as-is if we somehow end up empty
            display = raw_str

        display_cols.append(display)
        mapping[raw_str] = display

    return display_cols, mapping


# -------------------------
# Safe arithmetic evaluator
# -------------------------

def _eval_arith_node(node: ast.AST, env: Dict[str, Any]) -> float | None:
    """
    Evaluate a restricted arithmetic AST node.

    Allowed:
      - Names (column refs)
      - Constants (ints/floats)
      - +, -, *, /, %, unary -
      - Parentheses
    """
    if isinstance(node, ast.Expression):
        return _eval_arith_node(node.body, env)

    # Binary operations
    if isinstance(node, ast.BinOp):
        left = _eval_arith_node(node.left, env)
        right = _eval_arith_node(node.right, env)
        if left is None or right is None:
            return None

        if isinstance(node.op, ast.Add):
            return left + right
        if isinstance(node.op, ast.Sub):
            return left - right
        if isinstance(node.op, ast.Mult):
            return left * right
        if isinstance(node.op, ast.Div):
            try:
                return left / right
            except ZeroDivisionError:
                return None
        if isinstance(node.op, ast.Mod):
            try:
                return left % right
            except ZeroDivisionError:
                return None
        raise ValueError(f"Unsupported binary operator: {type(node.op).__name__}")

    # Unary -x
    if isinstance(node, ast.UnaryOp) and isinstance(node.op, ast.USub):
        val = _eval_arith_node(node.operand, env)
        return -val if val is not None else None

    # Constants
    if isinstance(node, ast.Constant) and isinstance(node.value, (int, float)):
        return float(node.value)

    # Name → env lookup
    if isinstance(node, ast.Name):
        raw_val = env.get(node.id)
        if raw_val is None:
            return None
        try:
            return float(raw_val)
        except (TypeError, ValueError):
            return None

    # Anything else is not allowed
    raise ValueError(f"Unsupported expression element: {type(node).__name__}")


def eval_arithmetic_expr(expr: str, env: Dict[str, Any]) -> float | None:
    """
    Safely evaluate a limited arithmetic expression like 'PRICE * VOLUME'.

    Returns float or None on error.
    """
    if not expr or not expr.strip():
        return None

    try:
        tree = ast.parse(expr, mode="eval")
        return _eval_arith_node(tree, env)
    except Exception:
        # On any parse/eval error we just return None
        return None


def build_report_context(
    *,
    title: str,
    question: str,
    sql: str,
    columns: List[str],
    rows: Sequence[Sequence],
    model: str,
    gen_ms: float | None = None,
    exec_ms: float | None = None,
    rows_est: int | str | None = None,
    generated_by: str = "AI Reporting System",
) -> dict:
    """
    Build a context dict that matches report.html placeholders.

    Also computes:
      - display_columns       (pretty headers, no underscores, Proper Case)
      - column_display_map    (raw -> display)
      - raw_columns           (for traceability / later customisation)
    """
    now = datetime.now()
    as_of_str = now.strftime("%Y-%m-%d %H:%M:%S")
    report_id = now.strftime("RPT-%Y%m%d-%H%M%S")
    print_date = now.strftime("%d/%m/%Y")  # Print Date
    print_time = now.strftime("%I:%M:%S %p").lower()  # Print Time

    # Normalise headers for display
    display_columns, column_display_map = normalize_display_columns(columns)

    ctx = {
        "title": title,
        "as_of_str": as_of_str,
        "generated_by": generated_by,
        "question": question,
        "model": model,
        "gen_ms": int(gen_ms) if gen_ms is not None else "-",
        "exec_ms": int(exec_ms) if exec_ms is not None else "-",
        "rows_est": rows_est if rows_est is not None else "-",

        # raw data
        "raw_columns": list(columns),      # keep original names
        "columns": list(columns),          # current raw columns used in table
        "rows": [list(r) for r in rows],
        "rows_returned": len(rows),

        # display metadata
        "display_columns": display_columns,
        "column_display_map": column_display_map,

        "sql": sql,
        "report_id": report_id,
        "generated_at": as_of_str,
        "print_date": print_date,
        "print_time": print_time,
    }

    # ------------- NEW FIELDS FOR THE FANCY TEMPLATE -------------

    # Header: user + print date/time + simple page number
    ctx["user_id"] = "ADMIN"          # override at caller if you know real user
    ctx["print_date"] = now.strftime("%d/%m/%Y")
    ctx["print_time"] = now.strftime("%I:%M:%S %p").lower()
    ctx["page_no"] = 1                # simplest: always "Page: 1"

    # From / To date and extra meta under the header
    # (caller can override these if it knows the filters)
    ctx.setdefault("from_date", "")
    ctx.setdefault("to_date", "")
    ctx.setdefault("extra_meta", [])

    # Numeric columns → right align automatically based on column names
    numeric_indexes: List[int] = []
    upper_cols = [str(c).upper() for c in columns]
    for idx, name in enumerate(upper_cols):
        if any(
            key in name
            for key in ("VOL", "VALUE", "AMOUNT", "PRICE", "CHARGE", "QTY", "RATE", "TOTAL")
        ):
            numeric_indexes.append(idx)
    ctx["numeric_indexes"] = numeric_indexes

    # Totals: try to compute common things if those columns exist
    def _sum_if_contains(keyword: str) -> float | None:
        try:
            col_idx = next(i for i, n in enumerate(upper_cols) if keyword in n)
        except StopIteration:
            return None
        total = 0.0
        for r in rows:
            if col_idx >= len(r):
                continue
            val = r[col_idx]
            if val is None:
                continue
            try:
                total += float(val)
            except (TypeError, ValueError):
                continue
        return total


    # Hide the technical/debug section by default in the PDF
    ctx.setdefault("show_technical", False)

    return ctx


# -------------------------
# Column customization
# -------------------------

def apply_column_customizations(context: dict, rules: dict | None) -> dict:
    """
    Apply user-defined column customisations to the report context.

    rules shape:
    {
      "include": { "FLAG_CODE": false, "FLAG_DESCRIPTION": false, ... },
      "rename": { "CLIENT_NAME": "Client", ... },
      "computed": [
        {
          "name": "TRADED_VALUE",
          "label": "Traded Value",          # optional
          "expression": "ORDER_PRICE * ORDER_VOLUME"
        },
        ...
      ]
    }
    """
    if not rules:
        return context

    include_map: Dict[str, bool] = {
        str(k): bool(v) for k, v in (rules.get("include") or {}).items()
    }
    rename_map: Dict[str, str] = {
        str(k): str(v) for k, v in (rules.get("rename") or {}).items() if v
    }
    computed_specs = rules.get("computed") or []

    # Start from current raw columns + rows
    cols: List[str] = list(context.get("columns") or [])
    rows: List[List[Any]] = [list(r) for r in context.get("rows") or []]

    # 1) Filter columns by include_map
    if include_map and cols:
        keep_indices: List[int] = []
        new_cols: List[str] = []
        for idx, c in enumerate(cols):
            if include_map.get(c, True):  # default = True
                keep_indices.append(idx)
                new_cols.append(c)

        if keep_indices:
            cols = new_cols
            rows = [[row[i] for i in keep_indices] for row in rows]

    # 2) Add computed columns (based on current cols)
    for spec in computed_specs:
        name = (spec.get("name") or "").strip()
        expr = (spec.get("expression") or "").strip()
        if not name or not expr:
            continue

        # Add optional label into rename map if provided
        label = (spec.get("label") or "").strip()
        if label:
            rename_map[name] = label

        base_cols = cols[:]   # columns available for env lookup
        cols.append(name)

        for row in rows:
            env = {}
            for i, col_name in enumerate(base_cols):
                if i < len(row):
                    env[col_name] = row[i]
            val = eval_arithmetic_expr(expr, env)
            row.append(val)

    # 3) Rebuild display columns (normalised) then apply renames
    norm_display, norm_map = normalize_display_columns(cols)
    display_cols: List[str] = []
    for raw, default_disp in zip(cols, norm_display):
        disp = rename_map.get(raw, default_disp)
        display_cols.append(disp)

    context["columns"] = cols
    context["rows"] = rows
    context["display_columns"] = display_cols
    context["column_display_map"] = {raw: disp for raw, disp in zip(cols, display_cols)}
    context["rows_returned"] = len(rows)
    context["customization_rules"] = rules

    return context


def apply_layout_grouping(context: dict, layout: dict | None) -> dict:
    """
    Apply layout / grouping information to the context.

    This does NOT change context['columns'] or context['rows'].
    It only adds:
      - layout_type: "flat" | "grouped"
      - group_by: list of column names actually used for grouping
      - detail_columns: columns to display inside each group table
      - groups: list of { "key": {...}, "rows": [...] }

    layout shape (from API):
    {
      "layout_type": "flat" | "grouped",
      "group_by": ["BROKER_NAME", ...],
      "detail_columns": ["BROKER_NAME", "EMAIL", ...]
    }
    """
    # Default: flat layout, no grouping
    layout = layout or {}
    layout_type = (layout.get("layout_type") or "flat").lower()

    columns: list[str] = list(context.get("columns") or [])
    rows: list[list[Any]] = [list(r) for r in context.get("rows") or []]

    # Always set some sane defaults on the context
    if "detail_columns" not in context:
        context["detail_columns"] = columns
    if "group_by" not in context:
        context["group_by"] = []
    if "groups" not in context:
        context["groups"] = []
    if "layout_type" not in context:
        context["layout_type"] = "flat"

    # If not grouped → just ensure layout_type="flat" and return
    if layout_type != "grouped":
        context["layout_type"] = "flat"
        return context

    group_by_cfg = layout.get("group_by") or []
    if not group_by_cfg or not columns or not rows:
        # Nothing to group by → fall back to flat layout
        context["layout_type"] = "flat"
        return context

    # Map column name -> index
    index_map: Dict[str, int] = {str(name): idx for idx, name in enumerate(columns)}

    # Keep only group_by columns that actually exist in the result set
    group_by: list[str] = [g for g in group_by_cfg if g in index_map]
    if not group_by:
        context["layout_type"] = "flat"
        return context

    # Build groups
    groups_map: Dict[tuple, Dict[str, Any]] = {}
    for row in rows:
        key_vals: list[Any] = []
        for col in group_by:
            idx = index_map[col]
            val = row[idx] if idx < len(row) else None
            key_vals.append(val)
        key_tuple = tuple(key_vals)

        if key_tuple not in groups_map:
            key_obj = {col: key_vals[i] for i, col in enumerate(group_by)}
            groups_map[key_tuple] = {"key": key_obj, "rows": []}
        groups_map[key_tuple]["rows"].append(row)

    groups = list(groups_map.values())

    # Decide which columns to show inside each group table
    detail_cfg = layout.get("detail_columns") or []
    if detail_cfg:
        detail_columns = [c for c in detail_cfg if c in columns]
        if not detail_columns:
            detail_columns = columns
    else:
        detail_columns = columns

    context["layout_type"] = "grouped"
    context["group_by"] = group_by
    context["detail_columns"] = detail_columns
    context["groups"] = groups

    return context


# -------------------------
# Renderers
# -------------------------

def build_html_report(context: dict) -> str:
    """Render the HTML report."""
    return render_report_html(context)


def build_csv_bytes(
    columns: List[str],
    rows: Sequence[Sequence],
    *,
    normalize_headers: bool = True,
) -> bytes:
    """
    Create CSV bytes from columns + rows (UTF-8 encoded).

    By default the header row uses normalised display columns
    (no underscores, Proper Case). Set normalize_headers=False if you
    explicitly want raw column names.
    """
    if normalize_headers:
        header_cols, _ = normalize_display_columns(columns)
    else:
        header_cols = columns

    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(header_cols)
    for r in rows:
        writer.writerow(r)
    return buf.getvalue().encode("utf-8")


def build_docx_bytes(context: dict) -> bytes:
    """
    Build a simple DOCX report:
    - Title
    - Question + meta
    - Table of results (with display headers)
    - SQL at the end
    """
    doc = Document()

    # Title
    doc.add_heading(context.get("title", "Report"), level=1)

    # Meta section
    p = doc.add_paragraph()
    p.add_run("Question: ").bold = True
    p.add_run(str(context.get("question", "")))
    p.add_run("\nModel: ").bold = True
    p.add_run(str(context.get("model", "")))
    p.add_run("\nGenerated at: ").bold = True
    p.add_run(str(context.get("generated_at", "")))

    # Table
    cols = context.get("display_columns") or context.get("columns", [])
    rows = context.get("rows", [])
    if cols:
        table = doc.add_table(rows=1, cols=len(cols))
        hdr_cells = table.rows[0].cells
        for i, c in enumerate(cols):
            hdr_cells[i].text = str(c)

        for row in rows:
            row_cells = table.add_row().cells
            for i, cell in enumerate(row):
                row_cells[i].text = "" if cell is None else str(cell)

    # SQL section
    doc.add_heading("SQL", level=2)
    doc.add_paragraph(str(context.get("sql", "")))

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
